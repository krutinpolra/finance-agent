"""
Turning an arbitrary uploaded file into a transactions table.

A CSV is already structured data, so pandas reads it directly. Everything
else - a photographed receipt, a scanned PDF statement, a Word doc - isn't,
so instead of writing a parser per file format, a single Claude call reads
the file directly (Claude accepts PDFs and images natively as message
content) and returns the same date/description/amount shape the rest of the
pipeline already expects. One extraction step, several input formats.
"""

import base64
import io
import json
import re

import pandas as pd
from docx import Document

SYSTEM_PROMPT = """You extract bank/credit-card transactions from a document
or image of a financial statement. Find every transaction line and return
ONLY a JSON array on one line, nothing else, in this exact shape:
[{"date": "YYYY-MM-DD", "description": "...", "amount": -12.34}, ...]

Use a negative amount for money going out (purchases, payments, fees) and a
positive amount for money coming in (deposits, refunds, income). If the year
isn't shown, infer it from context or use the most recent plausible year. If
you can't find any transactions, return an empty array: []
"""


def _parse_transactions_json(text: str) -> list[dict]:
    match = re.search(r"\[.*\]", text, re.DOTALL)
    if not match:
        return []
    try:
        return json.loads(match.group(0))
    except json.JSONDecodeError:
        return []


def _ask_claude_to_extract(client, content_block: dict) -> pd.DataFrame:
    response = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{
            "role": "user",
            "content": [content_block, {"type": "text", "text": "Extract every transaction from this."}],
        }],
    )
    final_text = next(b.text for b in response.content if b.type == "text")
    transactions = _parse_transactions_json(final_text)
    return pd.DataFrame(transactions, columns=["date", "description", "amount"])


def _docx_to_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def extract_transactions(client, uploaded_file) -> pd.DataFrame:
    """Takes a Streamlit UploadedFile and returns a DataFrame with
    date/description/amount columns, whatever format the source file was in.
    """
    name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if name.endswith(".csv"):
        return pd.read_csv(io.BytesIO(file_bytes))

    if name.endswith(".pdf"):
        block = {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": base64.b64encode(file_bytes).decode(),
            },
        }
        return _ask_claude_to_extract(client, block)

    if name.endswith((".png", ".jpg", ".jpeg")):
        media_type = "image/png" if name.endswith(".png") else "image/jpeg"
        block = {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": media_type,
                "data": base64.b64encode(file_bytes).decode(),
            },
        }
        return _ask_claude_to_extract(client, block)

    if name.endswith(".docx"):
        block = {"type": "text", "text": _docx_to_text(file_bytes)}
        return _ask_claude_to_extract(client, block)

    raise ValueError(f"Unsupported file type: {uploaded_file.name}")
